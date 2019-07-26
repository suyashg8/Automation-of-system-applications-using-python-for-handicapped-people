from pyautogui import *
import serial
from os import *
from threading import Thread
from time import *
import docx
import webbrowser
from PIL import Image
import pyautogui
from googlesearch import search
import cv2
import pyttsx


ser3=serial.Serial("COM3",9600)
finish=False
print("\t\tHello guys! Welcome to the voice operated system.")
print()

engine = pyttsx.init()
engine.say("hello guys! welcome to the voice operated system. ")
engine.runAndWait()
print("Here are some of the voice controlled commands:")
print()
print("1) Enable mouse control")
print("2) Open web browser")
print("3) Open presentation")
print("4) take a picture")
print("5) take a screenshot")
print("6) read a word file")
print("7) create a word file")

me="Me: "
system="System: "
nextline="\n"
mouse_controls_list = ['enable mouse usage','enable mouse control','start mouse usage','start mouse control', 'enable the mouse control', 'enable a mouse control','start the mouse control']
presentation_list=['open a presentation', 'open presentation', 'open the presentation','launch presentation']
search_list =['i want to search' , 'start search', 'choose search','select search','search']
screenshot_list=['take a screenshot','take screenshot', 'take the screenshot', 'click a screenshot','click screenshot','click the screenshot']
website_list=['open a website', 'open website','launch website','open the website','launch the website','launch a website']
browser_list=['open web browser','open browser','open a browser','open the browser','open the web browser','open a web browser']
read_word=['read a word file','read the word file','read word file']
take_picture=['take picture','take a picture','take the picture','click picture','click a picture']
create_word=['create a word file','create word file','create the word file','make a word file','make word file']
app_window=['change application window','change app window','switch app window','switch application window']
closed=['close application','exit','close','exit application','exit app','close app']
browser_tabs=['change tab','next tab','switch tab']
new_tabs=['open a new tab','new tab','move to new tab','open new tab']
change_tabs=['next tab','change tab','move to next tab','swithch tab']
create_w=['create a word file','create word file','create the word file','make a word file','make word file']



def word_file():
    print(system+"Enter the name of the documnet"+"\n")
    t=ser3.readline().strip().decode('utf-8')
    t=t.replace("*",'')
    t=t.replace("#",'')
    name=t        
    print(me+t)
    doc=docx.Document()
    doc.save(name+'.docx')
    print(system+"Launching the word file "+name+".docx\n")
    engine.say("Launching microsoft word")
    engine.runAndWait()
    startfile(name+".docx")
    word_typing()
    pyautogui.keyDown('ctrl')
    pyautogui.press('s')
    pyautogui.keyUp('ctrl')
    sleep(4)
    engine.say("Closing file")
    engine.runAndWait()
    close()


def mouse_control():
    print(system+"mouse control enabled"+"\n")
    engine.say("mouse control enabled")
    engine.runAndWait()
    
    Finish=True
    while Finish:
        t=ser3.readline().strip().decode('utf-8')
        t=t.replace("*",'')
        t=t.replace("#",'')
        print(me+t)
        words=t.split()
        if t=="left click":
            click()
        if t=="double click":
            doubleClick()
        if t=="right click":
            rightClick()
        if t=="done":
            Finish=False

        if t=="up":
            pyautogui.moveRel(0, -120, duration = 1)
        if t=="down":
            pyautogui.moveRel(0, 120, duration = 1)
        if t=="right":
            pyautogui.moveRel(120,0, duration = 1)
        if t=="left":
            pyautogui.moveRel(-120,0, duration = 1)
        if t=="slightly up":
            pyautogui.moveRel(0, -30, duration = 1)
        if t=="slightly down":
            pyautogui.moveRel(0, 30, duration = 1)
        if t=="slightly right":
            pyautogui.moveRel(30,0, duration = 1)
        if t=="slightly left":
            pyautogui.moveRel(-30,0, duration = 1)
        if t=="right diagonal up":
            pyautogui.moveRel(150, -150, duration = 1)
        if t=="right diagonal down":
            pyautogui.moveRel(150, 150, duration = 1)
        if t=="left diagonal up":
            pyautogui.moveRel(-150,150, duration = 1)
        if t=="left diagoanl down":
            pyautogui.moveRel(-150,0, duration = 1)
        if t=="top":
            pyautogui.moveTo(700,0, duration = 1)
        if t=="bottom":
            pyautogui.moveTo(700, 1050, duration = 1)
        if t=="extreme left":
            pyautogui.moveTo(0,400, duration = 1)
        if t=="extreme right":
            pyautogui.moveRel(1500,500, duration = 1)
        if t=="top left corner":
            pyautogui.moveTo(0,0, duration = 1)
        if t=="bottom left corner":
            pyautogui.moveTo(0,800 , duration = 1)
        if t=="top right corner":
            pyautogui.moveTo(1800,0, duration = 1)
        if t=="bottom right corner":
            pyautogui.moveTo(1800,800 , duration = 1)


    print(system+"Mouse control disabled"+"\n")
    engine.say("mouse control disabled")
    engine.runAndWait()

def new_tab():
    pyautogui.keyDown('ctrl')
    pyautogui.press('t')
    pyautogui.keyUp('ctrl')
    

    
def change_app_window():
    pyautogui.keyDown('alt')
    pyautogui.press('tab')
    Finish=True
    while Finish:
        t=ser3.readline().strip().decode('utf-8')
        t=t.replace("*",'')
        t=t.replace("#",'')
        print(me+t)
        if (t.lower()=="next"):
            pyautogui.press('tab')        
        if (t.lower()=="select"):
            pyautogui.keyUp('alt')
            Finish=False
        else:
            print(system+"cannot uderstand"+"\n")
            engine.say("cannot understand")
            engine.runAndWait()


def change_tab():
    pyautogui.keyDown('ctrl')
    pyautogui.press('tab')
    Finish=True
    while Finish:
        t=ser3.readline().strip().decode('utf-8')
        t=t.replace("*",'')
        t=t.replace("#",'')
        print(me+t)
        if (t.lower()=="next"):
            pyautogui.press('tab')        
        if (t.lower()=="select"):
            pyautogui.keyUp('ctrl')
            Finish=False
            engine.say("tab selected")
            engine.runAndWait()
        else:
            print("cannot uderstand"+"\n")
            engine.say("cannot understand")
            engine.runAndWait()


def typing():
        finish=False
        engine.say("start speaking text")
        engine.runAndWait()
        while not finish:
            t=ser3.readline().strip().decode('utf-8')
            t=t.replace("*",'')
            t=t.replace("#",'')
            print(me+t)
            n=t.lower()
            
            if(t=="stop typing"):
                finish=True
            if t.lower()=="press enter":
                finish=True
                pyautogui.press('enter')    
            else:
                tem1=t
                tem1=tem1.replace("fullstop",".")
                tem1=tem1.replace("comma",",")
                tem1=tem1.replace("question mark","?")
                tem1=tem1.replace("tab","   ")
                typewrite(tem1,interval=0.001)

def word_typing():
        finish=False
        engine.say("start speaking text")
        engine.runAndWait()
        while not finish:
            t=ser3.readline().strip().decode('utf-8')
            t=t.replace("*",'')
            t=t.replace("#",'')
            print(me+t)
            n=t.lower()
            
            if(t=="stop typing"):
                finish=True
                n=""
            if t.lower()=="next line":
                pyautogui.press('enter')
                n=""

            if n=="bold":
                pyautogui.keyDown('ctrl')
                pyautogui.press('b')
                pyautogui.keyUp('ctrl')
                n=""
            if n=="underline":
                pyautogui.keyDown('ctrl')
                pyautogui.press('u')
                pyautogui.keyUp('ctrl')
                n=""
            if n=="undo":
                pyautogui.keyDown('ctrl')
                pyautogui.press('z')
                pyautogui.keyUp('ctrl')
                n=""
            if n=="increase font":
                pyautogui.keyDown('ctrl')
                pyautogui.keyDown('shift')
                pyautogui.press('>')
                pyautogui.keyUp('shift')
                pyautogui.keyUp('ctrl')
                n=""
            if n=="backspace":
                pyautogui.keyDown('ctrl')
                pyautogui.press('backspace')
                pyautogui.keyUp('ctrl')
                n=""   
            if n=="decrease font":
                pyautogui.keyDown('ctrl')
                pyautogui.keyDown('shift')
                pyautogui.press('<')
                pyautogui.keyUp('shift')
                pyautogui.keyUp('ctrl')
                n=""
            
            else:
                tem1=n
                tem1=tem1.replace("full stop",".")
                tem1=tem1.replace("comma",",")
                tem1=tem1.replace("question mark","?")
                tem1=tem1.replace("tab","   ")
                typewrite(tem1,interval=0.001) 
    
def close():
    hotkey("alt","f4")

    
def scrolldown():
    for i in range (0,8):
        pyautogui.press('down')


def next_link():
     for i in range (0,11):
         pyautogui.press("tab")
     link=True
     
     while link:
         t=ser3.readline().strip().decode('utf-8')
         t=t.replace("*",'')
         t=t.replace("#",'')
         print(me+t)
         name=t.lower()
         if(t.lower()=="next link"):
             pyautogui.press("tab")
         if t.lower()=="done":
             link=False
         if t.lower()=="launch link":
             pyautogui.press("enter")
         if name in mouse_controls_list:
             mouse_control()
    
def scrollup():
    for i in range (0,8):
        pyautogui.press('up')
            
def screenshot():
    im = pyautogui.screenshot()
    print(system+"assign a name for the screenshot"+"\n")
    engine.say("assign a name for the screenshot")
    engine.runAndWait()

    t=ser3.readline().strip().decode('utf-8')
    t=t.replace("*",'')
    t=t.replace("#",'')
    print(me+t)
    name=t
    im.save(t+'.png')
    print(system+"Do you want to open the screenshot file?"+"\n")
    print(system+"Yes or No?")
    t=ser3.readline().strip().decode('utf-8')
    t=t.replace("*",'')
    t=t.replace("#",'')
    print(t)
    if(t.lower()=="yes"):
        try:  
            engine.say("opening the screenshot file")
            engine.runAndWait()

            startfile(name+".png")  
        except IOError: 
            print("File does not exist..")
    if(t.lower()=="no"):
        print("File saved")
    print("File saved"+"\n")
    sleep(10)
    close()
    
def scroll():
    pyautogui.scroll(1000)


while True:
    t=ser3.readline().strip().decode('utf-8')
    t=t.replace("*",'')
    t=t.replace("#",'')
    print(me+t)
    t=t.lower()
    command=""
    command=t
    if command in create_w:
        word_file()

    if command in screenshot_list:
        screenshot()
    
    if command in presentation_list:
            print(system+"Enter the name of the presentation"+"\n")
            t=ser3.readline().strip().decode('utf-8')
            t=t.replace("*",'')
            t=t.replace("#",'')
            name=t    
            startfile(name+".pptx")
            engine.say("opening presentation")
            engine.runAndWait()

            sleep(6);
            finish2=False
            hotkey("fn","f5")
            while not finish2:
                t=ser3.readline().strip().decode('utf-8')
                t=t.replace("*",'')
                t=t.replace("#",'')
                if t.lower()=="next":
                    typewrite("\n",interval=0.001)
                elif t.lower()=="back":
                    typewrite("\b",interval=0.001)
                elif t.lower()=="close presentation":
                    hotkey("alt","f4")
                    hotkey("alt","f4")
                    finish2=True
                    engine.say("closing")
                    engine.runAndWait()

                    print(system+"Presentation ended"+"\n")
    
    if command in take_picture:
        name=""
        camera = cv2.VideoCapture(0)
        return_value, image = camera.read()
        print(system+"assign a name for the webcam picture"+"\n")
        engine.say("assign a name for the webcam picture")
        engine.runAndWait()

        t=ser3.readline().strip().decode('utf-8')
        t=t.replace("*",'')
        t=t.replace("#",'')
        print(me+t)
        name=t
        cv2.imwrite(name+'.png', image)
        print(system+"Do you want to open the screenshot file?"+"\n")
        print(system+"Yes or No?")
        t=ser3.readline().strip().decode('utf-8')
        t=t.replace("*",'')
        t=t.replace("#",'')
        print(t)
        if(t.lower()=="yes"):
            try:  
                engine.say("opening the photo")
                engine.runAndWait()

                startfile(name+".png")  
            except IOError: 
                print("File does not exist..")
        if(t.lower()=="no"):
            print(system+"The photo is saved"+"\n")
        engine.say("The photo is saved")
        engine.runAndWait()
        
        del(camera)
        sleep(5)
        close()    

    if command in closed:
        close()
        
    
    
    if command in browser_list:
        print(system+"Select what do you want to do:")
        engine.say("Select what do you want to do")
        engine.runAndWait()
    
        print(system+"Search   or   open a website"+"\n")
        
        t=ser3.readline().strip().decode('utf-8')
        t=t.replace("*",'')
        t=t.replace("#",'')
        print(me+t)
        command=""
        command=t.lower()
        if(command in search_list):
            engine.say("Search selected")
            engine.runAndWait()
    
            t=ser3.readline().strip().decode('utf-8')
            t=t.replace("*",'')
            t=t.replace("#",'')
            print(me+t)
            name=t        
            google =name
            engine.say("Opening web browser")
            engine.runAndWait()
    
            webbrowser.open_new_tab('http://www.google.com/search?btnG=1&q=%s' % google)
            sleep(10)
            looking=True
            while looking:
                t=ser3.readline().strip().decode('utf-8')
                t=t.replace("*",'')
                t=t.replace("#",'')
                name=t.lower()
                if( t.lower()=="scroll up"):
                    scrollup()
                if name in change_tabs:
                    change_tab()
                if(t=="scroll down"):
                    scrolldown()
                if(name in new_tabs):
                    new_tab()
                if name=="start typing":
                    typing()
                if(t=="done"):
                    looking=False
                    hotkey("alt","f4")
                if name in mouse_controls_list:
                    mouse_control()                    
                if(t=="next link"):
                   next_link()
            p=1
            close()
                    
            
        if(command in website_list):
            engine.say("You want to open a website")
            engine.runAndWait()
            print(system+"Enter the name of the website excluding www."+"\n")
            engine.say("mention the website")
            engine.runAndWait()
            t=ser3.readline().strip().decode('utf-8')
            t=t.replace("*",'')
            t=t.replace("#",'')
            web=t
            engine.say("Opening web browser")
            engine.runAndWait()
    
            webbrowser.open_new_tab('http://www.'+web)
            sleep(10)
            looking=True
            while looking:
                t=ser3.readline().strip().decode('utf-8')
                t=t.replace("*",'')
                t=t.replace("#",'')
                name=""
                name=t.lower()
                if(name in new_tabs):
                    new_tab()
                
                if( t=="scroll up"):
                    scrollup()
                if(t=="scroll down"):
                    scrolldown()
                if name=="start typing":
                    typing()
                if name in change_tabs:
                    change_tab()
                if(t=="done"):
                    looking=False
                    hotkey("alt","f4")
                if name in mouse_controls_list:
                    mouse_control()
            
        
    if command in read_word:
        print(system+"Enter the name of the documnet"+"\n")
        t=ser3.readline().strip().decode('utf-8')
        t=t.replace("*",'')
        t=t.replace("#",'')
        name=t
        print(me+t)
        doc = docx.Document(name+'.docx')
        l=len(doc.paragraphs)
        for i in range(0,l):
            print(doc.paragraphs[i].text)
            engine.say(doc.paragraphs[i].text)
            engine.runAndWait()
    
        print("")
        print("End of File"+"\n")
        engine.say("End Of File")
        engine.runAndWait()
    

        
    if command in mouse_controls_list:
            mouse_control()        
            
    if command in app_window:
            change_app_window()
        
    if t.lower()=="start typing":
        typing()

    if t.lower()=="close application":
        engine.say("Closing application.")
        engine.runAndWait()
        hotkey("alt","f4")
    if t.lower()=="goodbye":
        break
    
        
    finish=False

hotkey("alt","c")
hotkey("alt","f4")
